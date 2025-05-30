from flask import jsonify
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
import uuid
from openpyxl import load_workbook
from tempfile import gettempdir
import sys


def scrape_data(mnv, template):
    if not mnv.isnumeric():
        raise ValueError("Mã nhân viên không hợp lệ.")

    options = webdriver.ChromeOptions()
    options.add_argument('--headless=new')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-extensions')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--blink-settings=imagesEnabled=false')

    browser = None
    try:
        service = Service(ChromeDriverManager().install())
        browser = webdriver.Chrome(service=service, options=options)

        browser.get("http://scfp.vn/Productscan.aspx")
        browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_txtMANV").send_keys(mnv)
        browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_rdbLoai_2").click()
        browser.find_element(By.ID, "ctl00_ContentPlaceHolder1_txtXem").click()

        WebDriverWait(browser, 5).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="ctl00_ContentPlaceHolder1_gvPd"]/tbody'))
        )
        tbody = browser.find_element(By.XPATH, '//*[@id="ctl00_ContentPlaceHolder1_gvPd"]/tbody')
        rows = tbody.find_elements(By.TAG_NAME, "tr")
        data = [[cell.text for cell in row.find_elements(By.TAG_NAME, 'td')] for row in rows]

        df = pd.DataFrame(data[:], columns=data[0])

        # Logic for phieu_xuat.docx
        if template == "phieu_xuat":
            df.iloc[:, 0] = pd.to_datetime(df.iloc[:, 0], format="%m/%d/%Y %I:%M:%S %p").dt.strftime('%d/%m/%Y')
            df.iloc[:, 3] = df.iloc[:, 4] + ' / ' + df.iloc[:, 3]
            df.iloc[:, 4] = df.iloc[:, 5]
            df.iloc[:, 5] = ""
            df = df.iloc[:, :-1]

        # Logic for KPH.xlsx
        elif template == "KPH":
            df["Ngày phát hiện"] = pd.to_datetime(df.iloc[:, 0], format="%m/%d/%Y %I:%M:%S %p").dt.strftime('%d/%m/%Y')
            df["Mã sản phẩm"] = df.iloc[:, 3]
            df["Tên sản phẩm"] = df.iloc[:, 4]
            df["Nhà cung cấp"] = df.iloc[:, 2]
            df["Đơn vị tính"] = ""
            df["Số lượng"] = df.iloc[:, 5]
            df["Nơi phát hiện"] = ""
            df["Tình trạng SP KPH"] = ""
            df["Nguyên nhân (lỗi KPH)(nếu có)"] = ""
            df["Người giao SP KPH (ghi rõ tên)"] = ""
            df["Đề nghị xử lý"] = ""

            df = df[
                [
                    "Ngày phát hiện",
                    "Mã sản phẩm",
                    "Tên sản phẩm",
                    "Nhà cung cấp",
                    "Đơn vị tính",
                    "Số lượng",
                    "Nơi phát hiện",
                    "Tình trạng SP KPH",
                    "Nguyên nhân (lỗi KPH)(nếu có)",
                    "Người giao SP KPH (ghi rõ tên)",
                    "Đề nghị xử lý"
                ]
            ]
        else:
            raise ValueError("Unknown template specified.")

        return df

    finally:
        if browser:
            browser.quit()
            
def generate_word(data):
    from datetime import datetime
    today_str = datetime.today().strftime("InHangRaQuay_%d%m%Y.docx")
    output_path = os.path.join(gettempdir(), today_str)

    template_path = resource_path(os.path.join("template", "phieu_xuat.docx"))
    document = Document(template_path)
    table = document.tables[0]

    # Fill in the table
    for i, row in enumerate(data, start=1):
        if i >= len(table.rows):
            table.add_row()
        for j, key in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(row[key])
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # Remove trailing empty rows
    for i in range(len(table.rows) - 1, 0, -1):
        if not any(cell.text.strip() for cell in table.rows[i].cells):
            table._element.remove(table.rows[i]._element)

    document.save(output_path)
    return output_path


def generate_excel(data):
    template_path = resource_path(os.path.join("template", "KPH.xlsx"))
    workbook = load_workbook(template_path)
    sheet = workbook.active

    start_row = 6  # Start filling from row 6
    for i, row_data in enumerate(data):
        row_index = start_row + i
        for j, key in enumerate([
            "Ngày phát hiện",
            "Mã sản phẩm",
            "Tên sản phẩm",
            "Nhà cung cấp",
            "Đơn vị tính",
            "Số lượng",
            "Nơi phát hiện",
            "Tình trạng SP KPH",
            "Nguyên nhân (lỗi KPH)(nếu có)",
            "Người giao SP KPH (ghi rõ tên)",
            "Đề nghị xử lý"
        ]):
            value = row_data.get(key, "")
            sheet.cell(row=row_index, column=j+1, value=value)

    output_filename = f"KPH_output_{uuid.uuid4().hex[:8]}.xlsx"
    output_path = os.path.join(gettempdir(), output_filename)
    workbook.save(output_path)

    return output_path

def resource_path(filename):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, filename)
    return os.path.join(os.path.abspath("."), filename)
